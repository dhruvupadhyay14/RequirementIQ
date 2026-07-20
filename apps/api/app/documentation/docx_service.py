from io import BytesIO
from docx import Document as WordDocument


class DOCXService:
    def export(self, title: str, content: str) -> bytes:
        document = WordDocument(); document.add_heading(title, 0)
        for line in content.splitlines():
            if line.startswith("# "): document.add_heading(line[2:], 0)
            elif line.startswith("## "): document.add_heading(line[3:], 1)
            elif line.startswith("- "): document.add_paragraph(line[2:], style="List Bullet")
            elif line: document.add_paragraph(line)
        stream = BytesIO(); document.save(stream); return stream.getvalue()
